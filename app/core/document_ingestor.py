
from typing import BinaryIO, Dict, Any
from datetime import datetime
import pdfplumber
from pdfminer.high_level import extract_text
import json
from docx import Document
from PIL import Image
import pytesseract
import pdf2image


class DocumentIngestor:
    def __init__(self, store_client, embedding_generator):
        self.store_client = store_client
        self.embedding_generator = embedding_generator

    def process_document(self, file: BinaryIO, filename: str, doc_id: str):
        """
        Process the document and ingest it into the database
        """

        file_type = filename.split('.')[-1].lower()
        metadata = self._extract_metadata(file, file_type)
        content = self._extract_content(file, file_type)
        if file_type == 'json':
            # For JSON files, parse and store as JSON string
            try:
                if isinstance(content, str):
                    json_content = json.loads(content)
                else:
                    json_content = content

                # Convert JSON to string for storage
                json_str = json.dumps(json_content)

                # Generate embedding for the JSON content
                embedding = self.embedding_generator.generate(json_str)

                # Store the JSON data as string
                document = self.store_client.collections.get("Document")
                document.data.insert(
                    properties={
                        "content": json_str,
                        "json": json_str,  # Store as JSON string
                        "metadata": json.dumps({
                            "filename": filename,
                            "total_records": len(json_content) if isinstance(json_content, list) else 1,
                            **metadata
                        }),
                        "doc_id": doc_id,
                        "chunk_id": 0,
                        "file_type": file_type,
                    },
                    vector=embedding
                )
            except Exception as e:
                raise e
        else:
            # For non-JSON files, use the original chunking logic
            # Chunkify the content
            chunks = self._chunkify_content(content)

            # Generate embeddings for each chunk
            for idx, chunk in enumerate(chunks):
                embedding = self.embedding_generator.generate(chunk)
                document = self.store_client.collections.get("Document")

                document.data.insert(
                    properties={
                        "content": chunk,
                        "json": None,  # No JSON for non-JSON files
                        "metadata": json.dumps({
                            "filename": filename,
                            "total_chunks": len(chunks),
                            **metadata
                        }),
                        "doc_id": doc_id,
                        "chunk_id": idx,
                        "file_type": file_type,
                    },
                    vector=embedding
                )

    def _extract_metadata(self, file: BinaryIO, file_type: str) -> Dict[str, Any]:
        """
        Extract metadata from the document
        """

        metadata = {
            "file_type": file_type,
            "timestamp": datetime.now().isoformat(),
            "extraction_method": "unknown"
        }

        try:
            match file_type:
                case "pdf":
                    pos = file.tell()
                    file.seek(0)

                    try:
                        with pdfplumber.open(file) as pdf:
                            pdf_meta = pdf.metadata or {}
                            metadata.update({
                                "title": pdf_meta.get("Title", "Untitled"),
                                "author": pdf_meta.get("Author", "Unknown"),
                                "page_count": len(pdf.pages),
                                "creation_date": pdf_meta.get("CreationDate", "Unknown"),
                                "modified_date": pdf_meta.get("ModDate", "Unknown"),
                            })
                    except Exception as e:
                        print(f"Error reading PDF metadata: {e}")
                        metadata.update({
                            "title": "Unknown",
                            "author": "Unknown",
                            "page_count": 0,
                            "error": str(e)
                        })

                    # Add extraction method to metadata
                    metadata["extraction_method"] = self.last_extraction_method or "unknown"
                    if self.last_extraction_method == "OCR":
                        metadata["ocr_processed"] = True

                    file.seek(pos)

                case "docx":
                    pos = file.tell()
                    file.seek(0)

                    doc = Document(file)
                    meta = doc.core_properties
                    metadata.update({
                        "file_type": file_type,
                        "title": meta.title,
                        "author": meta.author,
                        "creation_date": meta.created.isoformat() if meta.created else None,
                        "last_modified_by": meta.last_modified_by,
                    })

                    file.seek(pos)

                case "json":
                    content = file.read()
                    file.seek(0)

                    json_data = json.loads(content)
                    metadata.update({
                        "keys_count": len(json_data) if isinstance(json_data, dict) else len(json_data) if isinstance(json_data, list) else 0,
                        "data_type": type(json_data).__name__
                    })

                case "txt":
                    content = file.read().decode('utf-8')
                    file.seek(0)  # Reset file pointer
                    lines = content.split('\n')
                    metadata.update({
                        "line_count": len(lines),
                        "character_count": len(content),
                        "word_count": len(content.split())
                    })
        except Exception as e:
            print(f"Error extracting metadata: {e}")

        return {k: v for k, v in metadata.items() if v is not None}

    def _extract_content(self, file: BinaryIO, file_type: str) -> str:
        """
        Extract content from the document
        """

        match file_type:
            case "pdf":
                self.last_extraction_method = None  # Reset for each extraction
                text_content = ""

                # Validate PDF first
                try:
                    # Try to read as PDF first
                    file.seek(0)
                    with pdfplumber.open(file) as pdf:
                        text_content = "\n".join(
                            page.extract_text() or "" for page in pdf.pages)
                except Exception as e:
                    print(f"PDF text extraction failed: {e}")
                    text_content = ""

                # If no text is found or text is very short, try OCR
                if not text_content or len(text_content.strip()) < 50:
                    try:
                        file.seek(0)
                        images = pdf2image.convert_from_bytes(
                            file.read(),
                            fmt='jpeg',
                            dpi=300
                        )

                        text_contents = []
                        for image in images:
                            text = pytesseract.image_to_string(image)
                            if text.strip():  # Only add non-empty text
                                text_contents.append(text)

                        if text_contents:
                            text_content = "\n".join(text_contents)
                            self.last_extraction_method = "OCR"
                        else:
                            self.last_extraction_method = "FAILED"
                            text_content = "No text could be extracted from this document."

                    except Exception as e:
                        print(f"OCR extraction failed: {e}")
                        self.last_extraction_method = "FAILED"
                        text_content = "Failed to process document."
                else:
                    self.last_extraction_method = "text_only"

                file.seek(0)  # Reset file pointer
                return text_content

            case "docx":
                doc = Document(file)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])

            case "json":
                return json.dumps(json.load(file), indent=2)

            case "txt":
                return file.read().decode('utf-8')

            case _:
                raise ValueError(f"Unsupported file type: {file_type}")

    def _chunkify_content(self, content: str, chunk_size: int = 1000) -> list[str]:
        """
        Chunkify the content into smaller chunks
        """

        chunks = []
        for i in range(0, len(content), chunk_size):
            chunks.append(content[i:i+chunk_size])

        return chunks
